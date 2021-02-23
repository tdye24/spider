import time
import json
from selenium import webdriver
from selenium.webdriver import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import RGBColor


# 1 企业基本情况
# 1.1 企业概况
# 1.1.1 基本信息
def A(browser, query):
    print("调用A函数，爬取1.1.1基本信息")
    DATA = {}
    browser.find_element_by_id('searchDetail').send_keys(query)
    browser.find_element_by_id('searchDetail').send_keys(Keys.ENTER)
    print(f"查询：{query}")
    browser.switch_to.window(browser.window_handles[1])
    print("页面loading")
    time.sleep(3)

    # 当前融资
    try:
        res = browser.find_element_by_xpath(
            "//tbody/tr/td[2]/div/div/div[2]/div[@class='AgencyDiv']/span[contains(text("
            "), '轮')]").text
    except:
        res = ""
    finally:
        DATA['当前融资'] = res
        print(f'当前融资: {res}')

    link = browser.find_element_by_xpath('//tbody/tr/td[2]/div/div/div[2]/span[1]')
    link.click()
    print(f"点击链接，页面跳转...")
    browser.switch_to.window(browser.window_handles[-1])
    print("页面loading")
    time.sleep(3)

    print("抓取基本信息...")

    # 硬科技评级指数
    res = browser.find_element_by_xpath("//div[@class='comdetailstop_right']"
                                        "//div[@class='middle_right']"
                                        "//div[@class='item'][1]"
                                        "//div[@class='itembottom']"
                                        "/span").text
    DATA['硬科技评级指数'] = res
    print(f"硬科技评级指数: {res}")

    # 专利量
    res = browser.find_element_by_xpath("//div[@class='comdetailstop_right']"
                                        "//div[@class='middle_right']"
                                        "//div[@class='item'][2]"
                                        "//div[@class='itembottom']"
                                        "/span").text
    DATA['专利量'] = res
    print(f"专利量: {res}")

    # 发明团队
    res = browser.find_element_by_xpath("//div[@class='comdetailstop_right']"
                                        "//div[@class='middle_right']"
                                        "//div[@class='item'][3]"
                                        "//div[@class='itembottom']"
                                        "/span").text
    DATA['发明团队'] = res
    print(f"发明团队: {res}")

    # 工商信息
    browser.find_element_by_class_name('panoramicinsight') \
        .find_element_by_class_name('hideboxcontainer') \
        .find_element_by_class_name('middlemenutop') \
        .find_element_by_tag_name('span').click()
    print("加载工商信息")
    time.sleep(3)
    table = browser.find_element_by_class_name(name='industrybox').find_element_by_tag_name('table')
    # 企业名称
    res = table.find_elements_by_tag_name('tr')[0].find_elements_by_tag_name('td')[1].text
    DATA['企业名称'] = res
    print(f"企业名称: {res}")

    # 法定代表人
    res = table.find_elements_by_tag_name('tr')[0].find_elements_by_tag_name('td')[3].text
    DATA['法定代表人'] = res
    print(f"法定代表人: {res}")

    # 成立日期
    res = table.find_elements_by_tag_name('tr')[1].find_elements_by_tag_name('td')[1].text
    print(f"成立日期: {res}")
    DATA['成立日期'] = res

    # 统一社会信用代码
    res = table.find_elements_by_tag_name('tr')[2].find_elements_by_tag_name('td')[3].text
    DATA['统一社会信用代码'] = res
    print(f"统一社会信用代码: {res}")

    # 登记机关
    res = table.find_elements_by_tag_name('tr')[3].find_elements_by_tag_name('td')[1].text
    print(f"登记机关: {res}")
    DATA['登记机关'] = res

    # 所属地区
    res = table.find_elements_by_tag_name('tr')[4].find_elements_by_tag_name('td')[1].text
    print(f"所属地区: {res}")
    DATA['所属地区'] = res

    # 注册资本
    res = table.find_elements_by_tag_name('tr')[5].find_elements_by_tag_name('td')[1].text
    print(f"注册资本: {res}")
    DATA['注册资本'] = res

    # 公司类型
    res = table.find_elements_by_tag_name('tr')[5].find_elements_by_tag_name('td')[3].text
    print(f"公司类型: {res}")
    DATA['公司类型'] = res

    # 登记住所
    res = table.find_elements_by_tag_name('tr')[6].find_elements_by_tag_name('td')[1].text
    print(f"登记住所: {res}")
    DATA['登记住所'] = res

    # 经营范围
    res = table.find_elements_by_tag_name('tr')[6].find_elements_by_tag_name('td')[3].text
    DATA['经营范围'] = res
    print(f"经营范围: {res}")

    # 关闭窗口
    print("关闭小窗口")
    time.sleep(1)
    browser.find_element_by_xpath("//div[@class='el-dialog__wrapper industydetails_dia']/div/div[1]/button").click()

    return DATA


# 1.1.2 行业赛道
def B(browser, query):
    DATA = {}

    return DATA


# 1.1.3 标签画像
def C(browser, query):
    DATA = {}

    return DATA


# 1.1.4 竞合全景
def D(browser, query):
    DATA = {}

    return DATA


# 1.2 专利基本信息
# 1.2.1 专利基础信息
def E(browser, query):
    DATA = {}
    browser.find_element_by_xpath("//span[contains(text(), '专利资产洞察')]").click()
    time.sleep(0.5)
    browser.find_element_by_xpath("//span[contains(text(), '查看全部专利')]").click()
    time.sleep(0.5)
    browser.switch_to.window(browser.window_handles[-1])
    print(browser.current_url)
    time.sleep(5)
    print("加载专利列表...")

    for i in range(100):
        js = f'document.getElementsByClassName("el-table__body-wrapper is-scrolling-left")[0].scrollTop={i*1000}'
        browser.execute_script(js)
        print("scroll...")
        time.sleep(2)
    table = browser.find_element_by_tag_name('tbody')
    tr_lst = table.find_elements_by_tag_name('tr')
    print("共有专利: ", len(tr_lst))
    DATA['专利申请总量'] = len(tr_lst)
    print(tr_lst)
    for tr in tr_lst:
        tmp = {}
        # 简单法律状态名称
        res = tr.find_element_by_xpath('//td[47]//p[1]').text
        tmp['简单法律状态名称'] = res
        # 专利类型-中文
        res = tr.find_element_by_xpath('//td[51]//p[1]').text
        tmp['专利类型-中文'] = res
    return DATA


def run(username, pwd, query):
    browser = webdriver.Chrome('chromedriver.exe')
    browser.maximize_window()  # 最大化
    browser.get(url='https://www.sixlens.com/#/')
    browser.find_element_by_class_name(name='username').find_element_by_tag_name('input').send_keys(username)
    browser.find_element_by_class_name(name='pwd').find_element_by_tag_name('input').send_keys(pwd)
    print("输入验证码")
    time.sleep(10)
    browser.find_element_by_class_name('submit_btn').click()
    print("登录成功")
    # 获取cookies
    cookies = browser.get_cookies()
    with open('cookies.txt', 'w') as f:
        f.write(json.dumps(cookies))
    # print(cookies)
    # print(type(cookies))
    # time.sleep(10)
    # with open('cookies.txt') as f:
    #     cookies = json.loads(f.read())
    #     for cook in cookies:
    #         browser.add_cookie(cook)
    #     browser.refresh()
    # 等待页面加载完
    time.sleep(10)
    doc = Document()
    # 文档基本样式
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 1.企业基本情况
    doc.add_heading('1.企业基本情况', level=1)
    # 1.1 企业概况
    doc.add_heading('1.1 企业概况', level=2)
    # 1.1.1 基本信息
    doc.add_heading('1.1.1 基本信息', level=3)

    data = A(browser, query)
    print(data)
    # data = B(browser, query)
    # print(data)
    data = E(browser, query)
    print(data)

    doc.save(f'{query}.docx')

    browser.close()


def write(path):
    doc = Document()

    doc.styles['Normal'].font.name = u'宋体'

    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    paragraph = doc.add_paragraph()

    run = paragraph.add_run(" 程序员有话说")

    font = run.font

    #  设置字体大小

    font.size = Pt(24)

    #  设置水平居中

    paragraph_format = paragraph.paragraph_format

    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ------添加一段话-------

    content = "这是一个最好的时代，也是一个最坏的时代。好的是众多程序员都加入通过文字表达自己想法的步伐，" \
 \
              "不好的是依然围着技术转，始终不敢释放自我。你若不信，请听听他们的对话。"

    paragraph = doc.add_paragraph(content)

    paragraph_format = paragraph.paragraph_format

    #  第一行左边缩进

    paragraph_format.first_line_indent = Inches(0.3)

    # -----添加一个小标题------

    paragraph = doc.add_paragraph()

    run = paragraph.add_run("同程序员聊天")

    paragraph_format = paragraph.paragraph_format

    #  段前

    paragraph_format.space_after = Pt(15)

    #  段后

    paragraph_format.space_before = Pt(2)

    #  字体加粗和下划线

    font = run.font

    font.bold = True

    # ----跟销售人员聊天-----

    content = "大家好，我是西门吹水，做销售的，现去研发部找程旭猿聊聊天。你好，程旭猿，在忙什么呢?\n研究Python技术当中。\n"

    paragraph = doc.add_paragraph(content)

    run = paragraph.add_run("什么是派森来的?麻烦介绍一下。")

    font = run.font

    font.underline = True

    #  插入表格和内容

    table = doc.add_table(rows=3, cols=2, style="Medium Grid 1 Accent 1")

    table.cell(0, 0).text = "Python"

    table.cell(0, 1).text = "跨平台编程语言"

    table.cell(1, 0).text = "跨平台"

    table.cell(1, 1).text = "Windows、macOsS、Ubuntu等"

    table.cell(2, 0).text = "用途"

    table.cell(2, 1).text = "人工智能、Web、桌面系统..."

    # -----设置字体颜色------

    doc.add_paragraph("西门吹水：我一句都没听懂，怎么办法呀?")

    paragraph = doc.add_paragraph()

    run = paragraph.add_run("程旭猿：给你一张图片，自己体会去。")

    font = run.font

    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # -----添加图片，设置图片大小------

    doc.add_picture(r"good.jpeg", width=Inches(6.25))

    # ------保存word文档到当前目录下-------

    doc.save(path)


if __name__ == '__main__':
    run(username='shyyjsdx01', pwd='Sixlens123', query='上海同臣环保有限公司')
    write(path='demo.docx')
